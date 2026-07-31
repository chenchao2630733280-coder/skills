#!/usr/bin/env python
from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
PC_HINTS = ("pc", "admin", "web", "后台", "管理端", "管理后台", "平台", "系统管理")
MOBILE_HINTS = ("mobile", "app", "mini", "小程序", "移动端", "手机", "家长端", "用户端", "我的")


@dataclass
class Screenshot:
    path: Path
    title: str
    target: str
    width: int
    height: int


def natural_key(path: Path) -> list[Any]:
    parts = re.split(r"(\d+)", path.as_posix().lower())
    return [int(p) if p.isdigit() else p for p in parts]


def load_json(path: Path | None) -> dict[str, Any]:
    if not path:
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def clean_title(path: Path) -> str:
    stem = path.stem
    stem = re.sub(r"^\d+[\s._-]*", "", stem)
    stem = re.sub(r"[_-]+", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem or path.stem


def classify_image(path: Path) -> Screenshot:
    with Image.open(path) as img:
        width, height = img.size
    name = path.as_posix().lower()
    if any(hint.lower() in name for hint in MOBILE_HINTS):
        target = "mobile"
    elif any(hint.lower() in name for hint in PC_HINTS):
        target = "pc-admin"
    elif height > width * 1.18:
        target = "mobile"
    elif width >= height * 1.18 or width >= 900:
        target = "pc-admin"
    else:
        target = "pc-admin"
    return Screenshot(path=path, title=clean_title(path), target=target, width=width, height=height)


def collect_screenshots(root: Path) -> list[Screenshot]:
    files = [p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in IMAGE_EXTS]
    return [classify_image(p) for p in sorted(files, key=natural_key)]


def resolve_image(image_value: str, metadata_path: Path | None, screenshots_root: Path) -> Path:
    raw = Path(image_value)
    candidates = []
    if raw.is_absolute():
        candidates.append(raw)
    if metadata_path:
        candidates.append(metadata_path.parent / raw)
    candidates.append(screenshots_root / raw)
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[-1]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in (
        ("Heading 1", 18, "1F4E79"),
        ("Heading 2", 15, "1F4E79"),
        ("Heading 3", 12, "2F5597"),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc: Document, title: str, version: str, company: str, month: str) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("产品说明书")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"V{version}版")
    r.font.size = Pt(14)

    for _ in range(10):
        doc.add_paragraph()
    for text in (company, month):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(14)
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_change_record(doc: Document, completion: str, publish: str) -> None:
    doc.add_heading("修改记录", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ("版本", "日期", "说明", "作者")
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    values = ("V1.0", publish or completion, "初始版本", "项目组")
    for i, value in enumerate(values):
        set_cell_text(table.rows[1].cells[i], value)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(text)


def endpoint_title(project: str, target: str, suffix: str | None) -> str:
    if suffix:
        return f"{project}({suffix})"
    return f"{project}({'管理后台' if target == 'pc-admin' else '移动端'})"


def default_purpose(project: str, target: str) -> str:
    if target == "pc-admin":
        return (
            f"{project}面向系统管理员、运营人员及业务管理人员，提供基础数据维护、业务流程审核、"
            "统计分析、权限配置和运营管理等功能，帮助用户通过信息化方式完成日常管理工作。"
        )
    return (
        f"{project}移动端面向终端用户，提供信息浏览、在线预约或报名、资料提交、进度查询、"
        "消息反馈和个人中心等功能，帮助用户在手机端便捷完成业务办理。"
    )


def runtime_text(target: str) -> tuple[str, str, str]:
    if target == "pc-admin":
        return (
            "服务器端：Linux操作系统，Nginx作为反向代理服务器，后端可采用Java/Spring Boot、MySQL、Redis等常见企业级技术栈部署。",
            "客户端：Windows 10+/macOS 10.15+操作系统，Chrome、Edge、Firefox等主流浏览器，推荐1920×1080及以上分辨率访问。",
            "网络环境：需具备稳定互联网或政企内网访问能力，建议使用HTTPS协议保障数据传输安全。",
        )
    return (
        "服务器端：Linux操作系统，后端通过RESTful API向移动端提供数据服务，数据库和缓存组件按业务规模弹性配置。",
        "客户端：支持iOS、Android主流手机系统，可运行于微信小程序、移动H5或App环境，适配主流手机屏幕尺寸。",
        "网络环境：需具备移动网络或无线网络访问能力，系统对弱网场景提供加载提示和操作反馈。",
    )


def architecture_text(project: str, target: str, override: str | None) -> str:
    if override:
        return override
    if target == "pc-admin":
        return (
            f"{project}管理后台采用前后端分离架构，前端通过浏览器访问，后端提供统一接口服务。"
            "系统按权限控制不同角色的数据范围和操作能力，业务数据统一存储于数据库，并通过日志、校验和备份机制保障运行安全。"
        )
    return (
        f"{project}移动端采用客户端与服务端接口交互的架构，页面通过接口获取业务数据并提交用户操作。"
        "服务端负责身份校验、业务规则处理、状态流转和数据存储，移动端负责信息展示、表单录入和操作反馈。"
    )


def add_ui_rules(doc: Document, target: str) -> None:
    doc.add_heading("界面设计", level=1)
    doc.add_heading("用户界面设计规则", level=2)
    doc.add_heading("字体", level=3)
    add_para(doc, "系统界面采用微软雅黑、PingFang SC、Helvetica Neue、Arial等易读字体，正文以14px左右为主，标题根据层级适当加粗放大。")
    doc.add_heading("风格", level=3)
    if target == "pc-admin":
        add_para(doc, "管理后台采用左侧菜单导航与右侧内容区域结合的后台布局风格，页面结构清晰，适合列表查询、表单维护、审核处理和统计分析等管理场景。")
    else:
        add_para(doc, "移动端采用卡片式、列表式和表单式页面组合，突出核心入口和操作路径，适合用户在手机屏幕上快速浏览和办理业务。")
    doc.add_heading("色系", level=3)
    add_para(doc, "界面以白色和浅灰色作为基础背景，蓝色作为主操作和选中状态色，绿色、橙色、红色分别用于成功、提醒和危险状态。")
    doc.add_heading("控件", level=3)
    doc.add_heading("尺寸", level=4)
    add_para(doc, "控件尺寸遵循平台习惯，按钮、输入框、筛选项和弹窗保持统一间距，确保信息密度与可读性平衡。")
    doc.add_heading("布局", level=4)
    if target == "pc-admin":
        add_para(doc, "列表页面通常采用“搜索区域、操作按钮、数据表格、分页组件”的布局结构，表单页面通过弹窗、抽屉或独立页面承载编辑流程。")
    else:
        add_para(doc, "移动端页面自上而下排列核心内容，常用入口通过宫格、卡片或底部标签栏组织，表单字段按业务顺序纵向展示。")
    doc.add_heading("交互", level=4)
    add_para(doc, "系统为按钮点击、表单提交、数据加载、操作成功、操作失败和危险操作提供明确反馈，关键操作需进行校验或二次确认。")


def module_default_description(title: str, target: str) -> list[str]:
    if target == "pc-admin":
        return [
            f"【{title}】管理员进入该功能页面后，可查看相关业务数据，并通过查询条件快速筛选目标记录。",
            "页面通常提供新增、编辑、删除、详情、导出或审核等操作入口，系统在提交前进行必要校验，并在操作完成后给出结果提示。",
        ]
    return [
        f"【{title}】用户进入该页面后，可查看与当前业务相关的信息内容，并按照页面提示完成操作。",
        "移动端页面强调操作路径清晰和反馈及时，用户提交信息后系统会进行校验，并展示处理结果或状态变化。",
    ]


def add_image(doc: Document, path: Path, caption: str, target: str) -> None:
    if not path.exists():
        add_para(doc, f"截图文件未找到：{path}")
        return
    max_width = Cm(15.2 if target == "pc-admin" else 8.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=max_width)
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].font.size = Pt(9)
    c.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def normalize_modules(
    target: str,
    screenshots: list[Screenshot],
    meta: dict[str, Any],
    metadata_path: Path | None,
    screenshots_root: Path,
) -> list[dict[str, Any]]:
    target_meta = meta.get("targets", {}).get(target, {})
    modules = target_meta.get("modules")
    if modules:
        normalized = []
        for module in modules:
            item = dict(module)
            if item.get("image"):
                item["image_path"] = resolve_image(item["image"], metadata_path, screenshots_root)
            normalized.append(item)
        return normalized
    return [
        {
            "title": shot.title,
            "image_path": shot.path,
            "description": module_default_description(shot.title, target),
        }
        for shot in screenshots
        if shot.target == target
    ]


def add_intro(doc: Document, project: str, target: str, meta: dict[str, Any]) -> None:
    completion = meta.get("dev_completion_date", "")
    publish = meta.get("publish_date", "")
    purpose = meta.get("purpose") or default_purpose(project, target)

    doc.add_heading("引言", level=1)
    add_change_record(doc, completion, publish)
    doc.add_heading("简述", level=2)
    if completion:
        add_para(doc, f"开发完成：{completion}")
    if publish:
        add_para(doc, f"发表日期：{publish}")
    add_para(doc, f"开发目的：{purpose}")
    add_para(doc, f"在本篇文档中，列举描述了“{project}”主要功能的设计目的、产品界面、使用方法和部分应用场景。")
    add_para(doc, purpose)
    if target == "pc-admin":
        add_para(doc, "屏幕宽高：管理后台适配PC端1920×1080分辨率，推荐使用Chrome、Edge等主流浏览器访问。")
    else:
        add_para(doc, "屏幕宽高：移动端适配主流手机屏幕尺寸，页面控件满足触屏点击和阅读习惯。")

    doc.add_heading("运行环境", level=2)
    for text in runtime_text(target):
        add_para(doc, text)

    doc.add_heading("系统架构", level=2)
    add_para(doc, architecture_text(project, target, meta.get("architecture")))


def add_feature_summary(doc: Document, target: str, meta: dict[str, Any], modules: list[dict[str, Any]]) -> None:
    doc.add_heading("功能摘要", level=1)
    target_meta = meta.get("targets", {}).get(target, {})
    features = target_meta.get("features") or [m.get("title", "功能模块") for m in modules[:12]]
    if target == "pc-admin":
        add_para(doc, "管理后台围绕业务配置、数据维护、审核处理、统计分析和系统管理等场景展开，主要功能包括：")
    else:
        add_para(doc, "移动端围绕用户浏览、在线办理、状态查询和个人服务等场景展开，主要功能包括：")
    for feature in features:
        doc.add_paragraph(str(feature), style="List Bullet")


def add_functions(doc: Document, target: str, modules: list[dict[str, Any]]) -> None:
    doc.add_heading("功能展示及说明", level=1)
    doc.add_heading("后台-管理端" if target == "pc-admin" else "移动端", level=2)
    if not modules:
        add_para(doc, "当前未识别到该端截图，请补充截图或在metadata.json中配置功能模块。")
        return
    for module in modules:
        title = str(module.get("title") or "功能页面")
        doc.add_heading(title, level=3)
        descriptions = module.get("description") or module_default_description(title, target)
        if isinstance(descriptions, str):
            descriptions = [descriptions]
        for desc in descriptions:
            add_para(doc, str(desc))
        image_path = module.get("image_path")
        if image_path:
            add_image(doc, Path(image_path), title, target)


def build_document(
    target: str,
    screenshots: list[Screenshot],
    meta: dict[str, Any],
    metadata_path: Path | None,
    screenshots_root: Path,
    output_dir: Path,
) -> Path:
    project = meta.get("project_name", "软件系统")
    company = meta.get("company", "公司名称")
    version = str(meta.get("version", "1.0"))
    month = meta.get("month") or f"{date.today().year}年 {date.today().month} 月"
    suffix = meta.get("targets", {}).get(target, {}).get("title_suffix")
    title = endpoint_title(project, target, suffix)
    modules = normalize_modules(target, screenshots, meta, metadata_path, screenshots_root)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    setup_styles(doc)
    add_cover(doc, title, version, company, month)
    add_intro(doc, project, target, meta)
    add_feature_summary(doc, target, meta, modules)
    add_ui_rules(doc, target)
    add_functions(doc, target, modules)

    safe_title = re.sub(r'[<>:"/\\\\|?*]+', "", title)
    out = output_dir / f"{safe_title}产品说明书.docx"
    doc.save(out)
    return out


def selected_targets(force_target: str, screenshots: list[Screenshot], meta: dict[str, Any]) -> list[str]:
    if force_target == "both":
        return ["pc-admin", "mobile"]
    if force_target in {"pc-admin", "mobile"}:
        return [force_target]
    targets = set()
    targets.update(s.target for s in screenshots)
    targets.update(k for k in meta.get("targets", {}) if k in {"pc-admin", "mobile"})
    return [t for t in ("pc-admin", "mobile") if t in targets] or ["pc-admin"]


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Chinese software copyright DOCX manuals from screenshots.")
    parser.add_argument("--screenshots", required=True, help="Screenshot folder.")
    parser.add_argument("--output-dir", required=True, help="Output folder.")
    parser.add_argument("--metadata", help="Optional metadata JSON.")
    parser.add_argument("--project-name", help="Software/project name.")
    parser.add_argument("--company", help="Company name.")
    parser.add_argument("--version", help="Version number, default 1.0.")
    parser.add_argument("--force-target", choices=["auto", "pc-admin", "mobile", "both"], default="auto")
    args = parser.parse_args()

    screenshots_root = Path(args.screenshots).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    metadata_path = Path(args.metadata).resolve() if args.metadata else None
    meta = load_json(metadata_path)
    if args.project_name:
        meta["project_name"] = args.project_name
    if args.company:
        meta["company"] = args.company
    if args.version:
        meta["version"] = args.version

    screenshots = collect_screenshots(screenshots_root)
    targets = selected_targets(args.force_target, screenshots, meta)
    outputs = [
        build_document(target, screenshots, meta, metadata_path, screenshots_root, output_dir)
        for target in targets
    ]
    print("Generated:")
    for out in outputs:
        print(out)


if __name__ == "__main__":
    main()
