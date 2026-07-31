#!/usr/bin/env python3
"""Build a Chinese operation manual DOCX from a structured JSON spec."""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    set_run_font(run)
    return para


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_heading(level=level)
    para.clear()
    run = para.add_run(text)
    set_run_font(run, size=max(12, 18 - level), bold=True)


def add_labeled_text(doc: Document, label: str, value: str | Iterable[str] | None) -> None:
    if not value:
        return
    if isinstance(value, str):
        text = value
    else:
        text = "、".join(str(v) for v in value if str(v).strip())
    para = doc.add_paragraph()
    label_run = para.add_run(f"{label}：")
    set_run_font(label_run, bold=True)
    body_run = para.add_run(text)
    set_run_font(body_run)


def add_label(doc: Document, label: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(f"{label}：")
    set_run_font(run, bold=True)


def add_bullets(doc: Document, items: Iterable[Any], label: str | None = None) -> None:
    items = [item for item in items if item]
    if not items:
        return
    if label:
        add_label(doc, label)
    for item in items:
        text = item if isinstance(item, str) else item_to_text(item)
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(text)
        set_run_font(run)


def add_numbered(doc: Document, items: Iterable[Any], label: str | None = None) -> None:
    items = [item for item in items if item]
    if not items:
        return
    if label:
        add_label(doc, label)
    for item in items:
        text = item if isinstance(item, str) else item_to_text(item)
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(text)
        set_run_font(run)


def item_to_text(item: Any) -> str:
    if isinstance(item, dict):
        parts = []
        for key in ("name", "title", "actor", "action", "result", "description"):
            if item.get(key):
                parts.append(str(item[key]))
        return "；".join(parts)
    return str(item)


def resolve_image(path_text: str, spec_path: Path, screenshot_root: str | None) -> Path | None:
    raw = Path(path_text)
    candidates = []
    if raw.is_absolute():
        candidates.append(raw)
    if screenshot_root:
        candidates.append(Path(screenshot_root) / raw)
    candidates.append(spec_path.parent / raw)
    candidates.append(Path.cwd() / raw)
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def add_image(doc: Document, img_path: Path | None, caption: str, platform: str) -> None:
    if not img_path:
        add_labeled_text(doc, "截图", f"{caption or '未命名截图'}（图片文件未找到）")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    width = Cm(6.2) if "移动" in platform else Cm(15.2)
    run.add_picture(str(img_path), width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        set_run_font(cap_run, size=9)
        cap_run.italic = True


def normalize_list(value: Any) -> list[Any]:
    if not value:
        return []
    if isinstance(value, list):
        return value
    return [value]


def add_fields_table(doc: Document, fields: list[Any]) -> None:
    if not fields:
        return
    add_label(doc, "表单字段说明")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["字段", "必填", "说明"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
    for field in fields:
        row = table.add_row().cells
        if isinstance(field, dict):
            row[0].text = str(field.get("name", ""))
            row[1].text = "是" if field.get("required") else "否"
            row[2].text = str(field.get("description", ""))
        else:
            row[0].text = str(field)
            row[1].text = "待确认"
            row[2].text = ""
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)


def add_page(doc: Document, page: dict[str, Any], spec_path: Path, screenshot_root: str | None, platform: str, level: int = 3) -> None:
    add_heading(doc, str(page.get("title", "未命名页面")), level)
    add_labeled_text(doc, "入口路径", page.get("nav_path"))
    add_labeled_text(doc, "功能说明", page.get("purpose") or page.get("description"))

    screenshots = normalize_list(page.get("screenshot") or page.get("screenshots"))
    captions = normalize_list(page.get("caption") or page.get("captions"))
    for index, screenshot in enumerate(screenshots):
        caption = str(captions[index]) if index < len(captions) else str(page.get("title", "截图"))
        add_image(doc, resolve_image(str(screenshot), spec_path, screenshot_root), caption, platform)

    add_labeled_text(doc, "筛选条件", page.get("filters"))
    add_labeled_text(doc, "功能按钮", page.get("actions"))
    add_labeled_text(doc, "列表字段", page.get("columns"))
    add_fields_table(doc, normalize_list(page.get("fields")))
    add_numbered(doc, normalize_list(page.get("steps")), "操作步骤")
    add_bullets(doc, normalize_list(page.get("notes")), "注意")
    add_bullets(doc, normalize_list(page.get("warnings")), "风险提示")

    for subpage in normalize_list(page.get("subpages")):
        if isinstance(subpage, dict):
            add_page(doc, subpage, spec_path, screenshot_root, platform, min(level + 1, 5))


def add_workflows(doc: Document, workflows: list[dict[str, Any]], spec_path: Path, screenshot_root: str | None, platform: str) -> None:
    if not workflows:
        return
    add_heading(doc, "常见操作流程", 2)
    for idx, workflow in enumerate(workflows, 1):
        add_heading(doc, f"{idx}. {workflow.get('title', '未命名流程')}", 3)
        add_labeled_text(doc, "流程说明", workflow.get("summary"))
        for step in normalize_list(workflow.get("steps")):
            if not isinstance(step, dict):
                add_numbered(doc, [step])
                continue
            add_heading(doc, str(step.get("title", "步骤")), 4)
            add_labeled_text(doc, "操作角色", step.get("actor"))
            add_labeled_text(doc, "操作内容", step.get("action"))
            add_labeled_text(doc, "操作结果", step.get("result"))
            if step.get("screenshot"):
                add_image(
                    doc,
                    resolve_image(str(step["screenshot"]), spec_path, screenshot_root),
                    str(step.get("caption", step.get("title", ""))),
                    platform,
                )
        flow = normalize_list(workflow.get("flow"))
        if flow:
            add_labeled_text(doc, "流程概览", " -> ".join(str(item) for item in flow))


def add_faq(doc: Document, faq: list[dict[str, Any]]) -> None:
    if not faq:
        return
    add_heading(doc, "附录：常见问题", 2)
    for idx, item in enumerate(faq, 1):
        question = item.get("question") or item.get("q") or "问题"
        answer = item.get("answer") or item.get("a") or "待补充"
        add_heading(doc, f"{idx}. {question}", 3)
        add_paragraph(doc, str(answer))


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build(spec: dict[str, Any], spec_path: Path, output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = spec.get("title", "操作手册")
    platform = spec.get("platform", "")
    screenshot_root = spec.get("screenshot_root")

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(str(title))
    set_run_font(title_run, size=22, bold=True)
    if spec.get("subtitle"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(str(spec["subtitle"]))
        set_run_font(sub_run, size=14)

    doc.add_paragraph()
    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    meta_items = [
        ("适用端", platform or "待确认"),
        ("适用对象", spec.get("audience", "待确认")),
        ("版本", spec.get("version", "V1.0")),
        ("所属项目", spec.get("owner", "")),
        ("生成日期", spec.get("generated_date", date.today().isoformat())),
    ]
    for key, value in meta_items:
        if value:
            cells = meta.add_row().cells
            cells[0].text = key
            cells[1].text = str(value)
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run)

    doc.add_page_break()
    add_heading(doc, "目录", 2)
    for idx, section_spec in enumerate(normalize_list(spec.get("sections")), 1):
        if isinstance(section_spec, dict):
            add_paragraph(doc, f"{idx}. {section_spec.get('title', '未命名章节')}")
    if spec.get("workflows"):
        add_paragraph(doc, "常见操作流程")
    if spec.get("faq"):
        add_paragraph(doc, "附录：常见问题")

    for idx, section_spec in enumerate(normalize_list(spec.get("sections")), 1):
        if not isinstance(section_spec, dict):
            continue
        add_heading(doc, f"{idx}. {section_spec.get('title', '未命名章节')}", 2)
        add_labeled_text(doc, "章节说明", section_spec.get("description"))
        for page in normalize_list(section_spec.get("pages")):
            if isinstance(page, dict):
                add_page(doc, page, spec_path, screenshot_root, platform)

    add_workflows(doc, normalize_list(spec.get("workflows")), spec_path, screenshot_root, platform)
    add_faq(doc, normalize_list(spec.get("faq")))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a DOCX operation manual from manual_spec.json.")
    parser.add_argument("--spec", required=True, help="Path to manual_spec.json")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    spec_path = Path(args.spec).resolve()
    output_path = Path(args.output).resolve()
    with spec_path.open("r", encoding="utf-8") as f:
        spec = json.load(f)
    build(spec, spec_path, output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
